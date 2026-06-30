"""Extract text from PDF, DOCX, MD, TXT files."""
from __future__ import annotations
import logging
from typing import List, Dict, Any

logger = logging.getLogger(__name__)


def extract_pdf(path: str) -> List[Dict[str, Any]]:
    """Returns list of {"page": int, "text": str}. PyMuPDF primary, pypdf fallback."""
    try:
        import fitz  # PyMuPDF
        doc = fitz.open(path)
        pages = []
        for i, page in enumerate(doc):
            pages.append({"page": i + 1, "text": page.get_text("text")})
        doc.close()
        return pages
    except Exception as e:
        logger.warning("PyMuPDF failed (%s); falling back to pypdf", e)
        from pypdf import PdfReader
        reader = PdfReader(path)
        return [{"page": i + 1, "text": p.extract_text() or ""} for i, p in enumerate(reader.pages)]


def extract_docx(path: str) -> str:
    from docx import Document as DocxDocument
    doc = DocxDocument(path)
    parts = [p.text for p in doc.paragraphs if p.text and p.text.strip()]
    for table in doc.tables:
        for row in table.rows:
            parts.append(" | ".join(cell.text for cell in row.cells))
    return "\n\n".join(parts)


def extract_text_file(path: str) -> str:
    with open(path, "r", encoding="utf-8", errors="ignore") as f:
        return f.read()


def extract(path: str, file_type: str) -> Dict[str, Any]:
    """Returns {"pages": [{"page": int, "text": str}], "page_count": int|None}.
    For non-paginated formats, returns a single entry with page=None."""
    if file_type == "pdf":
        pages = extract_pdf(path)
        return {"pages": pages, "page_count": len(pages)}
    if file_type == "docx":
        return {"pages": [{"page": None, "text": extract_docx(path)}], "page_count": None}
    if file_type in ("markdown", "txt"):
        return {"pages": [{"page": None, "text": extract_text_file(path)}], "page_count": None}
    raise ValueError(f"Unsupported file_type: {file_type}")
